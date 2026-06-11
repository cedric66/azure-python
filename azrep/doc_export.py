"""Markdown to DOCX/PDF exporter with configurable document styling."""
import argparse
import copy
import os
import re
import sys
from html import escape
from pathlib import Path

DEFAULT_CONFIG = {
    "document": {
        "title": "",
        "author": "AKS Reporting Toolkit",
        "output_dir": "exports",
        "page": {
            "size": "A4",
            "margin_top_in": 0.75,
            "margin_bottom_in": 0.75,
            "margin_left_in": 0.8,
            "margin_right_in": 0.8,
        },
        "fonts": {
            "body": "Aptos",
            "heading": "Aptos Display",
            "code": "Consolas",
        },
        "sizes": {
            "body_pt": 10.5,
            "h1_pt": 20,
            "h2_pt": 15,
            "h3_pt": 12.5,
            "code_pt": 9,
        },
        "colors": {
            "heading": "1F4E78",
            "accent": "1F4E78",
            "table_header_fill": "1F4E78",
            "table_header_text": "FFFFFF",
            "code_fill": "F3F4F6",
            "body": "111827",
        },
        "paragraph": {
            "space_after_pt": 6,
            "line_spacing": 1.08,
        },
        "table": {
            "style": "Table Grid",
            "font_size_pt": 9,
        },
    }
}

HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")
BULLET_RE = re.compile(r"^\s*[-*+]\s+(.+)$")
ORDERED_RE = re.compile(r"^\s*\d+[.)]\s+(.+)$")
LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
TABLE_SEP_RE = re.compile(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$")


def _merge(a, b):
    out = copy.deepcopy(a)
    for k, v in (b or {}).items():
        if isinstance(v, dict) and isinstance(out.get(k), dict):
            out[k] = _merge(out[k], v)
        else:
            out[k] = v
    return out


def load_config(path=None):
    cfg = copy.deepcopy(DEFAULT_CONFIG)
    if not path:
        return cfg
    try:
        import yaml
    except ImportError:
        sys.exit("PyYAML is required for --config. Install requirements.txt in the container.")
    with open(path, "r", encoding="utf-8") as f:
        data = yaml.safe_load(f) or {}
    return _merge(cfg, data)


def clean_inline(text):
    text = LINK_RE.sub(lambda m: "%s (%s)" % (m.group(1), m.group(2)), text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = text.replace("**", "").replace("__", "").replace("*", "").replace("_", "")
    return text.strip()


def split_table_row(line):
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [clean_inline(c.strip()) for c in line.split("|")]


def parse_markdown(text):
    lines = text.splitlines()
    blocks, para = [], []
    i = 0

    def flush_para():
        nonlocal para
        if para:
            blocks.append({"type": "paragraph", "text": clean_inline(" ".join(para))})
            para = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            flush_para()
            i += 1
            continue
        if stripped.startswith("```"):
            flush_para()
            i += 1
            code = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            blocks.append({"type": "code", "text": "\n".join(code)})
            continue
        if "|" in stripped and i + 1 < len(lines) and TABLE_SEP_RE.match(lines[i + 1]):
            flush_para()
            header = split_table_row(stripped)
            i += 2
            rows = []
            while i < len(lines) and "|" in lines[i].strip() and lines[i].strip():
                rows.append(split_table_row(lines[i]))
                i += 1
            blocks.append({"type": "table", "header": header, "rows": rows})
            continue
        m = HEADING_RE.match(stripped)
        if m:
            flush_para()
            blocks.append({"type": "heading", "level": min(len(m.group(1)), 3),
                           "text": clean_inline(m.group(2))})
            i += 1
            continue
        bullet_items = []
        while i < len(lines):
            m = BULLET_RE.match(lines[i])
            if not m:
                break
            bullet_items.append(clean_inline(m.group(1)))
            i += 1
        if bullet_items:
            flush_para()
            blocks.append({"type": "bullets", "items": bullet_items})
            continue
        numbered_items = []
        while i < len(lines):
            m = ORDERED_RE.match(lines[i])
            if not m:
                break
            numbered_items.append(clean_inline(m.group(1)))
            i += 1
        if numbered_items:
            flush_para()
            blocks.append({"type": "numbers", "items": numbered_items})
            continue
        para.append(stripped)
        i += 1

    flush_para()
    return blocks


def _hex(value):
    return str(value or "000000").lstrip("#")


def export_docx(blocks, cfg, output_path, source_name):
    try:
        from docx import Document
        from docx.enum.text import WD_BREAK
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except ImportError:
        sys.exit("python-docx is required for DOCX export. Install requirements.txt.")

    doc_cfg = cfg["document"]
    doc = Document()
    sec = doc.sections[0]
    page = doc_cfg["page"]
    sec.top_margin = Inches(float(page["margin_top_in"]))
    sec.bottom_margin = Inches(float(page["margin_bottom_in"]))
    sec.left_margin = Inches(float(page["margin_left_in"]))
    sec.right_margin = Inches(float(page["margin_right_in"]))

    props = doc.core_properties
    props.author = str(doc_cfg.get("author") or "")
    props.title = str(doc_cfg.get("title") or Path(source_name).stem)

    colors = doc_cfg["colors"]
    fonts = doc_cfg["fonts"]
    sizes = doc_cfg["sizes"]
    normal = doc.styles["Normal"]
    normal.font.name = fonts["body"]
    normal.font.size = Pt(float(sizes["body_pt"]))
    normal.font.color.rgb = RGBColor.from_string(_hex(colors["body"]))
    normal.paragraph_format.space_after = Pt(float(doc_cfg["paragraph"]["space_after_pt"]))
    normal.paragraph_format.line_spacing = float(doc_cfg["paragraph"]["line_spacing"])

    for level, key in ((1, "h1_pt"), (2, "h2_pt"), (3, "h3_pt")):
        style = doc.styles["Heading %d" % level]
        style.font.name = fonts["heading"]
        style.font.size = Pt(float(sizes[key]))
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(_hex(colors["heading"]))
        style.paragraph_format.space_before = Pt(12 if level == 1 else 8)
        style.paragraph_format.space_after = Pt(6 if level == 1 else 4)

    title = doc_cfg.get("title") or Path(source_name).stem
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.name = fonts["heading"]
    run.font.size = Pt(float(sizes["h1_pt"]) + 2)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(_hex(colors["accent"]))
    p.paragraph_format.space_after = Pt(10)
    p.add_run().add_break(WD_BREAK.LINE)

    def shade(cell, fill):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), _hex(fill))
        tc_pr.append(shd)

    for b in blocks:
        if b["type"] == "heading":
            doc.add_paragraph(b["text"], style="Heading %d" % b["level"])
        elif b["type"] == "paragraph":
            doc.add_paragraph(b["text"])
        elif b["type"] == "bullets":
            for item in b["items"]:
                doc.add_paragraph(item, style="List Bullet")
        elif b["type"] == "numbers":
            for item in b["items"]:
                doc.add_paragraph(item, style="List Number")
        elif b["type"] == "code":
            p = doc.add_paragraph()
            r = p.add_run(b["text"])
            r.font.name = fonts["code"]
            r.font.size = Pt(float(sizes["code_pt"]))
        elif b["type"] == "table":
            width = max(len(b["header"]), max((len(r) for r in b["rows"]), default=0))
            table = doc.add_table(rows=1, cols=width)
            table.style = doc_cfg["table"].get("style") or "Table Grid"
            for j, val in enumerate(b["header"]):
                cell = table.rows[0].cells[j]
                cell.text = val
                shade(cell, colors["table_header_fill"])
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor.from_string(_hex(colors["table_header_text"]))
            for row in b["rows"]:
                cells = table.add_row().cells
                for j, val in enumerate(row[:width]):
                    cells[j].text = val
            doc.add_paragraph()

    doc.save(output_path)
    return output_path


def export_pdf(blocks, cfg, output_path, source_name):
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4, letter
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import (Preformatted, SimpleDocTemplate, Spacer,
                                        Table, TableStyle, Paragraph)
    except ImportError:
        sys.exit("reportlab is required for PDF export. Install requirements.txt.")

    doc_cfg = cfg["document"]
    page = doc_cfg["page"]
    page_size = letter if str(page.get("size", "")).lower() == "letter" else A4
    styles = getSampleStyleSheet()
    fonts = doc_cfg["fonts"]
    sizes = doc_cfg["sizes"]
    color_cfg = doc_cfg["colors"]

    body = ParagraphStyle("BodyCustom", parent=styles["BodyText"],
                          fontName="Helvetica", fontSize=float(sizes["body_pt"]),
                          leading=float(sizes["body_pt"]) * 1.25,
                          spaceAfter=float(doc_cfg["paragraph"]["space_after_pt"]))
    heading = {
        1: ParagraphStyle("H1Custom", parent=styles["Heading1"],
                          fontSize=float(sizes["h1_pt"]), textColor=colors.HexColor("#" + _hex(color_cfg["heading"]))),
        2: ParagraphStyle("H2Custom", parent=styles["Heading2"],
                          fontSize=float(sizes["h2_pt"]), textColor=colors.HexColor("#" + _hex(color_cfg["heading"]))),
        3: ParagraphStyle("H3Custom", parent=styles["Heading3"],
                          fontSize=float(sizes["h3_pt"]), textColor=colors.HexColor("#" + _hex(color_cfg["heading"]))),
    }
    code_style = ParagraphStyle("CodeCustom", parent=styles["Code"],
                                fontName="Courier", fontSize=float(sizes["code_pt"]),
                                backColor=colors.HexColor("#" + _hex(color_cfg["code_fill"])),
                                leftIndent=6, rightIndent=6, spaceAfter=8)

    story = []
    title = doc_cfg.get("title") or Path(source_name).stem
    story.append(Paragraph(escape(title), ParagraphStyle(
        "TitleCustom", parent=styles["Title"], fontSize=float(sizes["h1_pt"]) + 2,
        textColor=colors.HexColor("#" + _hex(color_cfg["accent"])))))
    story.append(Spacer(1, 0.15 * inch))

    for b in blocks:
        if b["type"] == "heading":
            story.append(Paragraph(escape(b["text"]), heading[b["level"]]))
        elif b["type"] == "paragraph":
            story.append(Paragraph(escape(b["text"]), body))
        elif b["type"] == "bullets":
            for item in b["items"]:
                story.append(Paragraph("&#8226; " + escape(item), body))
        elif b["type"] == "numbers":
            for idx, item in enumerate(b["items"], 1):
                story.append(Paragraph("%d. %s" % (idx, escape(item)), body))
        elif b["type"] == "code":
            story.append(Preformatted(b["text"], code_style))
        elif b["type"] == "table":
            data = [b["header"]] + b["rows"]
            tbl = Table(data, repeatRows=1)
            tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#" + _hex(color_cfg["table_header_fill"]))),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#" + _hex(color_cfg["table_header_text"]))),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), float(doc_cfg["table"]["font_size_pt"])),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#D1D5DB")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]))
            story.append(tbl)
            story.append(Spacer(1, 0.12 * inch))

    doc = SimpleDocTemplate(output_path, pagesize=page_size,
                            topMargin=float(page["margin_top_in"]) * inch,
                            bottomMargin=float(page["margin_bottom_in"]) * inch,
                            leftMargin=float(page["margin_left_in"]) * inch,
                            rightMargin=float(page["margin_right_in"]) * inch)
    doc.build(story)
    return output_path


def export_markdown(md_path, to="all", config=None, out_dir=None, output_name=None):
    cfg = load_config(config)
    md_path = Path(md_path)
    if not md_path.exists():
        sys.exit("Markdown file not found: %s" % md_path)
    text = md_path.read_text(encoding="utf-8")
    blocks = parse_markdown(text)
    doc_cfg = cfg["document"]
    out_dir = Path(out_dir or doc_cfg.get("output_dir") or "exports")
    out_dir.mkdir(parents=True, exist_ok=True)
    stem = output_name or md_path.stem

    outputs = []
    if to in ("docx", "all"):
        outputs.append(export_docx(blocks, cfg, str(out_dir / (stem + ".docx")), md_path.name))
    if to in ("pdf", "all"):
        outputs.append(export_pdf(blocks, cfg, str(out_dir / (stem + ".pdf")), md_path.name))
    return outputs


def main(argv=None):
    p = argparse.ArgumentParser(description="Convert Markdown to DOCX/PDF with style config",
                                formatter_class=argparse.ArgumentDefaultsHelpFormatter)
    p.add_argument("markdown", help="input .md file")
    p.add_argument("--to", choices=["docx", "pdf", "all"], default="all")
    p.add_argument("--config", help="YAML style config, e.g. report_style.example.yaml")
    p.add_argument("--out-dir", help="output directory")
    p.add_argument("--output-name", help="output basename without extension")
    args = p.parse_args(argv)
    outputs = export_markdown(args.markdown, args.to, args.config, args.out_dir, args.output_name)
    for path in outputs:
        print("Wrote: %s" % path)
    return outputs


if __name__ == "__main__":
    main()
